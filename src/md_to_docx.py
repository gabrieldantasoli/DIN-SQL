"""
md_to_docx.py — Converte markdown -> .docx (pandoc via pypandoc) e FORMATA as
tabelas (bordas em grade + cabecalho em negrito e sombreado), embutindo as figuras.

Uso: python src/md_to_docx.py ARTIGO.md ARTIGO.docx
"""

import os
import sys

import pypandoc
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)


def shade_cell(cell, fill="D9E1F2"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def format_tables(doc):
    for t in doc.tables:
        set_table_borders(t)
        for cell in t.rows[0].cells:          # cabecalho: negrito + sombreado
            shade_cell(cell)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                if not p.runs and p.text:
                    p.add_run(p.text).bold = True
    return len(doc.tables)


def main():
    md, out = sys.argv[1], sys.argv[2]
    workdir = os.path.dirname(os.path.abspath(md)) or "."
    os.chdir(workdir)
    md_name = os.path.basename(md)
    pypandoc.convert_file(
        md_name, "docx", outputfile=out, format="gfm",
        extra_args=["--resource-path=.", "--standalone", "-M", "lang=pt-BR"],
    )
    d = docx.Document(out)
    n = format_tables(d)
    d.save(out)
    print(f"{out} gerado — {n} tabelas formatadas, {len(d.tables)} no total")


if __name__ == "__main__":
    main()
